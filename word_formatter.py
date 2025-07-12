import os
import shutil
from pypdf import PdfReader
from docx2pdf import convert
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, RGBColor, Cm, Pt
from copy import deepcopy
from docx.oxml import OxmlElement

from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls

import time
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import logging
from queue import Queue
from threading import Thread
from PIL import Image

# Configure the logger
verb= 0
count = 0
logging.basicConfig(filename='error_log.txt', 
                    level=logging.ERROR, 
                    format='%(asctime)s %(levelname)s:%(message)s')
INVALID_FOLDER = 'invalid/'


def add_page_numbers(doc):
    """
    Add page numbers to each section's footer in the document.
    """
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Create a run for the page number
        run = paragraph.add_run()
        
        # Create the field for the page number
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        # Append the elements to the run
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(fldChar3)
        
        # Set the font for the run
        run.font.size = Pt(8)
        run.font.name = "Barlow"


def extract_module_name_from_specific_cell(doc):
    """
    Extracts the module name from a specific table and cell in the document, excluding quantity and version.
    
    Args:
        doc (Document): A Document object representing the Word document.
    
    Returns:
        str: The extracted module name or None if not found.
    """
    try:
        # Access Table 5, Row 1, Cell 1
        cell_text = doc.tables[5].rows[1].cells[1].text.strip()
        print(f"Extracted raw module name: {cell_text}")

        # Split the text to remove the quantity part (e.g., "9 x ") and other unwanted details
        parts = cell_text.split("x", 1)  # Split at the first 'x'
        if len(parts) > 1:
            # Extract only the module name, removing version and other extra parts
            module_name = parts[1].strip()
            module_name_parts = module_name.split()  # Split into components
            clean_module_name = " ".join(module_name_parts[:5])  # Combine first 4 parts for the module name
            print(f"Formatted module name: {clean_module_name}")
            return clean_module_name
        else:
            print("Module name format unexpected; unable to split.")
            return None

    except IndexError:
        print("Failed to extract module name from the specified table and cell.")
        return None
def update_module_name(doc):
    """
    Updates the module name on the cover page to match the one extracted from a specific cell.
    
    Args:
        doc (Document): A Document object representing the Word document.
    """
    # Extract the module name from the specified table and cell
    module_name = extract_module_name_from_specific_cell(doc)
    
    if not module_name:
        print("Module name not found; cannot update cover page.")
        return

    # Replace the module name on the cover page
    cover_page_found = False
    for para in doc.paragraphs:
        if "IBC MonoSol" in para.text:  # Assuming this is the placeholder on the cover
            cover_page_found = True
            # Debugging: Log before replacement
            print(f"Original paragraph: {para.text}")
            
            # Clear the existing text
            para.clear()
            
            # Add new run with H2 formatting
            run = para.add_run(module_name)
            run.bold = True
            run.font.size = Pt(20)  # Set font size for H2
            run.font.name = "Barlow"  # Example font name, adjust as needed
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set color to black
            run.font.color.rgb = RGBColor(250, 168, 32)

            # Set paragraph alignment to center (if needed)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            print(f"Updated paragraph with module name as H2: {para.text}")  # Debugging output
            break

    if not cover_page_found:
        print("Cover page module name placeholder not found.")


def darken_first_row_bottom_border(document):
    # Define the namespace URI directly in the attribute setting
    namespace_uri = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    for table in document.tables:
        if table.rows:  # Check if there are rows in the table
            first_row = table.rows[0]
            for cell in first_row.cells:
                # Accessing the cell's XML and ensuring it has tcBorders
                tc = cell._element
                tcBorders = tc.find('.//w:tcBorders', namespaces={'w': namespace_uri})
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tc.append(tcBorders)
                
                # Modify or add the bottom border to be darker and thicker
                bottom_border = tcBorders.find('.//w:top', namespaces={'w': namespace_uri})
                if bottom_border is None:
                    bottom_border = OxmlElement('w:top')
                    tcBorders.append(bottom_border)
                
                # Set the style of the border
                bottom_border.set(f'{{{namespace_uri}}}val', 'single')  # Style of the border
                bottom_border.set(f'{{{namespace_uri}}}sz', '5')       # Size of the border, making it thicker
                bottom_border.set(f'{{{namespace_uri}}}color', '000000')  # Color of the border, making it black

def convert_jp2_to_jpg(image_path):
    """
    Convert a .jp2 image to .jpg format with a white background.
    """
    # Load the jp2 image
    with Image.open(image_path) as img:
        # Create a new white background image
        white_background = Image.new("RGB", img.size, (255, 255, 255))
        
        # Paste the jp2 image onto the white background
        white_background.paste(img, mask=img.split()[3] if img.mode == 'RGBA' else None)
        
        # Save the new image as .jpg
        jpg_path = image_path.replace('.jp2', '.jpg')
        white_background.save(jpg_path, 'JPEG')
        print(f"Converted {image_path} to {jpg_path}")
        
    return jpg_path
def clear_folder_contents(file_name, folder_path):
    # Check if the folder exists
    if not os.path.exists(folder_path):
        print(f"Folder does not exist: {folder_path}")
        return
    
    # Iterate through the files and directories in the folder
    for file_or_dir in os.listdir(folder_path):
        path = os.path.join(folder_path, file_or_dir)
        try:
            if os.path.isfile(path):
                if path.endswith('.docx'):
                    if file_or_dir == file_name or file_or_dir.startswith('~$'):
                        print(f"Removing file: {path}")
                        os.remove(path)
                else:
                    print(f"Removing file: {path}")
                    os.remove(path)
            elif os.path.isdir(path):
                print(f"Removing directory and its contents: {path}")
                shutil.rmtree(path)
        except Exception as e:
            print(f"Failed to remove {path}. Reason: {e}")


def set_font_to_barlow(output_doc):
    # Iterate through all paragraphs in the document
    for paragraph in output_doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Barlow'
    
    # Iterate through all tables in the document
    for table in output_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Barlow'

        # Set font in headers and footers
    for section in output_doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Barlow'
                
        footer = section.footer
        for paragraph in footer.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Barlow'
def extract_raw_document_images(filepath):
    fileName = os.path.basename(filepath)
    folder_path = os.path.dirname(filepath)

    fileName = fileName.replace(".docx", "")
    convert(f"{folder_path}/{fileName}.docx", f"{folder_path}/{fileName}.pdf")
    pdfPath = f"{folder_path}/{fileName}.pdf"
    path = f"{folder_path}/images"
    
    if os.path.exists(path) and os.path.isdir(path):
        shutil.rmtree(path)

    os.makedirs(f"{folder_path}/images")

    reader = PdfReader(f"{folder_path}/{fileName}.pdf")
    page = reader.pages[2]
    count = 0
    images = []
    imageNames = []
    count = 0
    i = 0
    for page in reader.pages:
        for img in page.images:
            if i != 1 and img.name not in imageNames:
                count += 1
                images.append(img)
                imageNames.append(img.name)
                with open(f"{folder_path}/images/" + str(count) + '.' + img.name.split('.')[1],"wb") as fp:
                    fp.write(img.data)
        i+=1
            
def copy_paragraph(output_doc, paragraph):

    output_paragraph = output_doc.add_paragraph()
    # Alignment data of whole paragraph
    output_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
    i=0
    for row in paragraph.runs:
        output_row = output_paragraph.add_run(row.text)
        # Font data
        output_row.style.name = "Normal"
        # Size of font data
        if row.font.size != None:
            output_row.font.size = row.font.size-1000
        else:
            output_row.font.size = row.font.size-1000
        # Bold data
        output_row.bold = row.bold
        # Italic data
        output_row.italic = row.italic
        # Underline data
        output_row.underline = row.underline
        # Color data
        output_row.font.color.rgb = row.font.color.rgb

def add_h1(output_doc, text):
    heading = output_doc.add_heading(level=1)
    run = heading.add_run(text)
    run.bold = True
    run.font.name = 'Barlow'
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(250, 168, 32)
       
def add_h2(output_doc, text):
    heading = output_doc.add_heading(text, 2)
    heading.style.font.name = 'Barlow'
    heading.style.font.size = 200000
    heading.style.font.bold = False
    heading.line_spacing_rule = WD_LINE_SPACING.SINGLE
    heading.style.font.color.rgb = RGBColor(250, 168, 32)
    
def add_h3(output_doc, text):
    heading = output_doc.add_heading(text, 3)
    heading.style.font.name = 'Barlow'
    heading.style.font.size = 150000
    heading.style.font.bold = False
    heading.style.font.color.rgb = RGBColor(128, 128, 128)

def set_table_borders(table, color="E8E9EB"):
    """
    Set table borders to a specific color (hex code without #).
    """
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')

    # Define border styles with correct namespace handling
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f'w:{border_name}')
        border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')  # Border style
        border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '4')        # Border width
        border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '0')     # Border space
        border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', color)   # Border color
        tblBorders.append(border)

    tbl.tblPr.append(tblBorders)

def darken_title_line(table):
    """
    Darken the border line after the title row of the table.
    """
    first_row = table.rows[0]
    for cell in first_row.cells:
        cell_tc = cell._element
        tcBorders = cell_tc.find('.//w:tcBorders', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            cell_tc.append(tcBorders)

        # Create or update the bottom border
        bottom_border = tcBorders.find('w:bottom', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if bottom_border is None:
            bottom_border = OxmlElement('w:bottom')
            tcBorders.append(bottom_border)

        bottom_border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')  # Border style
        bottom_border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '12')       # Border width
        bottom_border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '0')     # Border space
        bottom_border.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '000000')# Border color (black)

def remove_vertical_borders(table):
    tbl = table._tbl
    for cell in tbl.iter(qn('w:tc')):
        tcPr = cell.tcPr
        tcBorders = tcPr.xpath('./w:tcBorders')
        if not tcBorders:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for border in ['left', 'right']:
            border_element = tcBorders.find(qn(f'w:{border}'))
            if border_element is None:
                border_element = OxmlElement(f'w:{border}')
                tcBorders.append(border_element)
            border_element.set(qn('w:val'), 'nil')

def format_table(output_doc):
    width = (Inches(4.5), Inches(4.5), Inches(1.5))
    t = 0
    for table in output_doc.tables:
        table.style = "Table Grid"
        set_table_borders(table, color="E8E9EB")  # Set the border color to gray
        darken_title_line(table)  # Darken the line after the title

        i = 0
        for r in table.rows:
            j=0
            if len(table.rows[i].cells) == 2:
                width = (Inches(6), Inches(6))
            elif len(table.rows[i].cells) == 3:
                width = (Inches(6.5), Inches(4.5), Inches(2))
            elif len(table.rows[i].cells) == 4:
                width = (Inches(4), Inches(3), Inches(1), Inches(4))
            elif len(table.rows[i].cells) == 5:
                width = (Inches(2), Inches(4), Inches(3), Inches(2), Inches(2))
            elif len(table.rows[i].cells) == 6:
                width = (Inches(2), Inches(2), Inches(2), Inches(2), Inches(2), Inches(2))
            elif len(table.rows[i].cells) == 7:
                width = (Inches(0.5), Inches(2), Inches(2), Inches(2), Inches(2.5), Inches(1.5), Inches(1.5))
            for cell in table.rows[i].cells:
                table.rows[i].cells[j].width = width[j]
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                j+=1
            i+=1
        t+=1
        remove_vertical_borders(table)
def add_cell_to_row(row):
    """
    Add a new cell to a row in a Word table by manipulating the underlying XML.
    """
    new_cell = OxmlElement('w:tc')
    row._tr.append(new_cell)
    return new_cell

def remove_empty_paragraphs(doc):
    """
    Remove empty paragraphs from the document while preserving those that contain images or other non-empty content.
    """
    # Define the namespace map for 'w'
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    for paragraph in doc.paragraphs:
        # Check if the paragraph has no text and no non-empty runs
        if not paragraph.text.strip() and all(not run.text.strip() for run in paragraph.runs):
            # Check if the paragraph contains any inline shapes (images) by looking for 'w:drawing'
            if not paragraph._element.findall('.//w:drawing', namespaces=nsmap):
                p = paragraph._element
                p.getparent().remove(p)


def remove_empty_sections(doc):
    """
    Remove empty sections from the document to avoid unnecessary empty pages.
    """
    for i, section in enumerate(doc.sections[:-1]):  # Skip the last section
        if not any(para.text.strip() for para in section.footer.paragraphs + section.header.paragraphs):
            # If both the header and footer are empty, remove the section break
            p = section._sectPr
            p.getparent().remove(p)
def set_column_widths(table, widths):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                cell.width = widths[idx]
                tcPr = cell._element.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(widths[idx].cm * 1440)))  # Convert inches to twips
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

                
def format_table_with_picture(output_doc, tableNo, imagePath):
    # Set the table style
    table = output_doc.tables[tableNo]
    table.style = "Table Grid"
    
    # Define desired widths: 4 inches for the first, 2 inches for the second, 2 inches for the third, and 4 inches for the last
    column_widths = [Inches(4), Inches(2), Inches(2), Inches(4)]
    set_column_widths(table, column_widths)
    
    # Ensure the table has at least 4 cells, adding if necessary
    first_row = table.rows[0]
    while len(first_row.cells) < 4:
        add_cell_to_row(first_row)  # Add new cell using XML manipulation
    
    # Access the desired cell (index 3, which is the fourth cell)
    target_cell = first_row.cells[3]
    
    # Clear existing paragraphs in the target cell
    for paragraph in target_cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    # Add a new paragraph and insert the image
    paragraph = target_cell.add_paragraph()
    run = paragraph.add_run()
    
    try:
        # Attempt to add the picture to the paragraph
        run.add_picture(imagePath, width=Cm(4.5), height=Cm(4.5))
    except Exception as e:
        print(f"Error adding image {imagePath}: {e}")
def copy_table(output_doc, table):
    p = output_doc.paragraphs[-1]
    new_tbl = deepcopy(table._tbl)
    p._p.addnext(new_tbl)

def title_run(r):
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(250, 168, 32)
    r.font.name = "Barlow (Heading)"
    
def replace_variables(output_doc, raw_doc):
    txbx = raw_doc.inline_shapes._body.xpath('//w:txbxContent')
    address_lines = []
    id = ''
    
    # Extract information from text boxes
    for tx_idx, tx in enumerate(txbx):
        children = tx.getchildren()
        for child_idx, child in enumerate(children):
            if child.text:
                if child.text.startswith("Angebotsnr."):
                    id = child.text
                else:
                    address_lines.extend(parse_address(child.text))
    
    # Extract module, kw, and date
    module = raw_doc.tables[1].cell(4, 1).paragraphs[0].text
    print(module)
    kw = raw_doc.tables[1].cell(2, 1).paragraphs[0].text
    date = raw_doc.paragraphs[0].text
    
    # Add additional address information
    additional_address = raw_doc.tables[0].cell(1, 0).paragraphs[0].text
    address_lines.extend(parse_address(additional_address))
    
    try:
        additional_address_line = raw_doc.tables[0].cell(1, 0).paragraphs[1].text
        address_lines.extend(parse_address(additional_address_line))
    except IndexError:
        pass

    # Use a set to track unique address lines
    unique_address_lines = list(dict.fromkeys(address_lines))  # Preserve order while removing duplicates

    # Remove address lines from the first page content if they are not needed
    # Assume paragraph indices 5-7 hold the address, adjust if necessary
    paragraphs_to_remove = [5, 6, 7]
    for i in sorted(paragraphs_to_remove, reverse=True):
        if i < len(output_doc.paragraphs):
            output_doc.paragraphs[i].clear()

    i = 0
    for p in output_doc.paragraphs:
        text = p.text
        if i == 2:
            p.text = ""
            r = p.add_run(text.replace("0.00", kw))
            title_run(r)
        if i == 3:
            p.text = ""
            r = p.add_run(text.replace("0", module))
            title_run(r)
        elif i == 5:
            p.text = ""
            if len(unique_address_lines) > 0:
                r = p.add_run(unique_address_lines[0])
                print(unique_address_lines[0])
                title_run(r)
        elif i == 6:
            p.text = ""
            if len(unique_address_lines) > 2:
                r = p.add_run(unique_address_lines[2])
                print(unique_address_lines[2])
                title_run(r)
        elif i == 7:
            p.text = ""
            if len(unique_address_lines) > 1:
                r = p.add_run(unique_address_lines[1])
                print(unique_address_lines[1])
                title_run(r)
        elif i == 8:
            p.text = ""
            r = p.add_run(date)
            title_run(r)
        i += 1

def set_font(paragraph, font_name):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        run._element.rPr.rFonts.set(qn('w:cs'), font_name)
def parse_address(address_text):
    """
    Parse the address into a list of lines based on whether it contains a comma.
    """
    if ',' in address_text:
        # Compact format with street and city in one line
        return [line.strip() for line in address_text.split(',')]
    else:
        # Standard format with multiple lines
        return address_text.splitlines()

def parse_address(address_text):
    """
    Parse the address into a list of lines. The last part of the address should keep the city and country together.
    """
    if ',' in address_text:
        # Split address based on commas
        parts = [part.strip() for part in address_text.split(',')]
        
        # Handle case where the last part might be the country
        if len(parts) > 2:
            # Combine the city and country into a single line
            combined = ', '.join(parts[-2:])
            return parts[:-2] + [combined]
        else:
            return parts
    else:
        # Standard format with multiple lines
        return address_text.splitlines()

def prepare_header(output_doc, raw_doc):
    txbx = raw_doc.inline_shapes._body.xpath('//w:txbxContent')
    address_lines = []
    date = ''
    id = ''

    # Extract information from text boxes
    for tx_idx, tx in enumerate(txbx):
        children = tx.getchildren()
        for child_idx, child in enumerate(children):
            if child.text:
                if child.text.startswith("Angebotsnr."):
                    id = child.text
                else:
                    # Parse and add the address lines
                    address_lines.extend(parse_address(child.text))

    # Extract date from paragraphs
    if raw_doc.paragraphs:
        date = raw_doc.paragraphs[0].text

    # Extract additional address information from tables if available
    try:
        additional_address = raw_doc.tables[0].cell(1, 0).paragraphs[0].text
        address_lines.extend(parse_address(additional_address))
    except IndexError:
        pass

    # Prepare the header section
    section = output_doc.sections[0]
    header = section.header
    section.different_first_page_header_footer = True
    t = header.add_table(2, 2, Inches(24))

    # Cell (0, 0) content
    cell_00 = t.cell(0, 0)
    cell_00.text = ''
    cell_00.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = cell_00.paragraphs[0]

    # Add parsed address lines without duplications and empty lines
    unique_addresses = set()
    for line in address_lines:
        line = line.strip()  # Remove leading/trailing whitespace
        if line and line not in unique_addresses:
            r = p.add_run(line)
            r.font.size = Pt(8)
            r.add_break()  # Only add a break if there's another line to follow
            unique_addresses.add(line)
    set_font(p, "Barlow")

    # Remove the last line break if it exists
    if p.runs and p.runs[-1].text.endswith('\n'):
        p.runs[-1].text = p.runs[-1].text.rstrip('\n')

    # Cell (1, 0) content
    cell_10 = t.cell(1, 0)
    cell_10.text = ''
    cell_10.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = cell_10.paragraphs[0]
    r11 = p.add_run(date)
    r11.font.size = Pt(8)
    r11.add_break()

    # Split `id` and handle cases where it doesn't contain a space
    id_parts = id.split(" ")
    if len(id_parts) > 1:
        r21 = p.add_run(id_parts[1])
    else:
        r21 = p.add_run(id)  # Fallback: use the entire `id` string
    r21.font.size = Pt(8)
    set_font(p, "Barlow")

    # Cell (0, 1) content
    cell_01 = t.cell(0, 1)
    cell_01.text = ''
    cell_01.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = cell_01.paragraphs[0]
    r21 = p.add_run()
    r21.add_picture("assets/template_images/header.png", Cm(1.5), Cm(1.25))

    # Cell (1, 1) content
    cell_11 = t.cell(1, 1)
    cell_11.text = ''
    cell_11.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = cell_11.paragraphs[0]
    r31 = p.add_run("Solardach 24 GmbH")
    r31.font.color.rgb = RGBColor(250, 168, 32)
    r31.font.size = Pt(8)
    r31.add_break()
    r42 = p.add_run("Sicher und zuverlässig")
    r42.font.color.rgb = RGBColor(250, 168, 32)
    r42.font.size = Pt(8)
    set_font(p, "Barlow")



def prepare_footer(output_doc):
    footer = output_doc.sections[0].footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("Solardach24 GmbH")
    r1.add_break()
    r2 = p.add_run("Reinacherstrasse 261 ∙ 4053 Basel")
    r2.add_break()
    r3 = p.add_run("Collègegasse 9 ∙ 2502 Biel/Bienne")
    r3.add_break()
    r4 = p.add_run("+41 61 511 22 22 ∙ office@solardach24.ch ∙ CHE-152.292-000")
    set_font(p, "Barlow")   

def add_toc(output_doc):
    add_h1(output_doc, "INHALTSVERZEICHNIS", )
    output_doc.add_paragraph(" ")
    paragraph = output_doc.add_paragraph()
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)
    run = paragraph.add_run()
    run.font.name = 'Barlow'
    run.font.size = Pt(15)
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-1" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar3 = OxmlElement('w:updateFields') 
    fldChar3.set(qn('w:val'), 'true') 
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    
    run.font.color.rgb = RGBColor(250, 168, 32)
    
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    p_element = paragraph._p
    
def add_picture_inline(output_doc, picture_path, width, height):
    """
    Adds a picture to the document if it exists.
    """
    # Check for .png, .jpg, and .jp2 extensions
    for extension in ['.png', '.jpg', '.jp2']:
        full_path = picture_path + extension
        if os.path.exists(full_path):
            if extension == '.jp2':
                # Convert jp2 to jpg
                full_path = convert_jp2_to_jpg(full_path)
            
            # Add the image to the document
            output_doc.add_picture(full_path, width, height)
            return  # Exit once the image is added successfully
        else:
            continue
    
    # Log an error if the image is not found
    return

def extract_para_style(raw, para_style):
    para = []
    for p in raw.paragraphs:
        if p.style.name.startswith(para_style):
                para.append(p.text)
    return para

def process_files(queue, template_path, output_folder):
    global flag
    while True:
        file_name, src_path = queue.get()
        if file_name[0]=='0':
            flag = 0
            print('filename starts with 0')
        else:
            flag = 1
            print('filename does not start with 0')
        try:
            main(file_name, src_path, template_path, output_folder)
        finally:
            queue.task_done()

# Watchdog event handler
class NewFileHandler(FileSystemEventHandler):
    def __init__(self, template_path, output_folder, queue):
        self.template_path = template_path
        self.output_folder = output_folder
        self.queue = queue
    
    def on_created(self, event):
        if event.is_directory:
            return
        if event.src_path.endswith('.docx'):
            file_name = os.path.basename(event.src_path)
            if not file_name.startswith('~$'):
                time.sleep(1)  # Ensure the file is fully written
                self.queue.put((file_name, event.src_path))

    def move_file_with_retry(self, src, dst, max_retries=5, delay=1):
        for _ in range(max_retries):
            try:
                shutil.move(src, dst)
                return
            except (PermissionError, FileNotFoundError):
                time.sleep(delay)
        print(f"Failed to move {src} to {dst} after {max_retries} retries.")

# Main function to set up watchdog observer
def set_(watch_folder, template_path, output_folder):
    queue = Queue()
    event_handler = NewFileHandler(template_path, output_folder, queue)
    
    observer = Observer()
    observer.schedule(event_handler, path=watch_folder, recursive=False)
    observer.start()

    # Start worker thread to process files from the queue
    worker_thread = Thread(target=process_files, args=(queue, template_path, output_folder), daemon=True)
    worker_thread.start()

    try:
        while True:
            time.sleep(0.02)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()
    queue.join()  # Wait for all tasks to be processed 


    # Save the updated document

# Function to remove the prefix from the title
def remove_prefix_from_title(doc):
    # Iterate through all paragraphs to find the title
    for para in doc.paragraphs:
        if para.text.startswith("Projektbericht - "):
            # Remove the prefix
            para.text = para.text.replace("Projektbericht - ", "", 1)
            break

def main(fileName, filepath, template_path, output_folder):
    global count, verb, flag
    global INVALID_FOLDER
    try:
        print(f'New document added: {fileName}')
        raw = Document(f'{filepath}')
        remove_prefix_from_title(raw)


        template = Document(f'{template_path}')
        doc = Document()
        print(raw)
        print(f"Total number of tables: {len(raw.tables)}")
        
        # Extract images from raw document - these a document specific images
        extract_raw_document_images(filepath)
        folder_path = os.path.dirname(filepath)
        doc.add_picture("assets/template_images/image1.png", width=Inches(6), height=Inches(4))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ## headings upto table of contents
        doc.add_paragraph(" ")
        for paragraph in template.paragraphs[12:20]:
            copy_paragraph(doc, paragraph)

        # -------------------------------------------------  
        doc.add_paragraph(" ")
        add_toc(doc)
        doc.add_page_break()
        # ---------------------------------------------------

        tableIndex = 1
        h1_index = 0
        h2_index = 0
        h3_index = 0
        pic_index = 2
        h1 = extract_para_style(raw, 'Heading 1')
        h2 = extract_para_style(raw, 'Heading 2')
        h3 = extract_para_style(raw, 'Heading 3')
        para = []
        for p in raw.paragraphs:
            if p.text != '':
                para.append(p.text)
        try:
            add_h1(doc, f"1. {h1[h1_index]}")
            h1_index+=1
        except:
            pass
        add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(6), height=Inches(4))
        pic_index+=1
        
        doc.add_paragraph(" ")
        if any("PV-Anlage" in h for h in h2):
            # This condition checks if "PV-Anlage" is in any of the headings of h2
            add_h2(doc, "PV-Anlage")
            
            # Search for the index of the paragraph containing "PV-Anlage"
            i = next((idx for idx, p in enumerate(para) if "PV-Anlage" in p), None)
            
            if i is not None:
                add_h3(doc, para[i+4])
                
                # Remove the processed 'PV-Anlage' from the list
                #para.remove('PV-Anlage')
                
                # Copy the table and add the picture
                copy_table(doc, raw.tables[tableIndex])
                print(f"PV-Anlage table copied - {tableIndex}")
                tableIndex += 1

                add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(6), height=Inches(4))
                
                # Increment the picture index
                pic_index += 1
        else:
            print("PV-Anlage not found in headings")
        
        if "Ertragsprognose" in h2 :
            add_h2(doc, "Ertragsprognose")
            add_h3(doc, "Ertragsprognose")
            copy_table(doc, raw.tables[tableIndex])
            print(f"Ertragsprognose table copied - {tableIndex}")
            # print(tableIndex)
            tableIndex+=1

        # ----------------------------------------
        doc.add_page_break()
        # ----------------------------------------
        try:
            add_h1(doc, f"2. {h1[h1_index]}")
            h1_index+=1
        except:
            pass
        if "Überblick" in h2 :
            add_h2(doc, "Überblick")

            add_h3(doc, "Anlagendaten")
            copy_table(doc, raw.tables[tableIndex])
            print(f"Anlagendaten table copied - {tableIndex}")
            # print(tableIndex)
            tableIndex+=1
            
            add_h3(doc, "Klimadaten")
            copy_table(doc, raw.tables[tableIndex])
            print(f"Klimadaten table copied - {tableIndex}")
            # print(tableIndex)
            tableIndex+=1
            
            add_h3(doc, "Verbrauch")
            copy_table(doc, raw.tables[tableIndex])
            print(f"Verbrauch table copied - {tableIndex}")
            # print(tableIndex)
            tableIndex+=1

            doc.add_paragraph("")
            add_picture_inline(doc, f"assets/images/{pic_index}", width=Inches(6), height=Inches(4))
            pic_index+=1

        if "Modulflächen" in h2 :
            doc.add_page_break()
            add_h2(doc, "Modulflächen")
            k = para.index("Modulflächen")
            
            for i in range(k+1, len(para), 3):
                if 'Modulfläche' in para[i]:
                    
                    add_h2(doc, para[i])
                    add_h3(doc, para[i+1])
                    copy_table(doc, raw.tables[tableIndex])
                    print(f"Modulfläche {i} table copied - {tableIndex}")
                    tableIndex+=1
                    doc.add_paragraph("")
                    add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(6), height=Inches(4))
                    pic_index+=1
                else:
                    break

        add_h2(doc, "Horizontlinie, 3D-Planung")
        add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(6), height=Inches(4))
        pic_index+=1

        # ----------------------------------------
        doc.add_page_break()
        # ----------------------------------------

        if "Wechselrichterverschaltung" in h2:
            add_h2(doc, "Wechselrichterverschaltung")
            
            # Find the index of "Wechselrichterverschaltung" in paragraphs
            i = para.index("Wechselrichterverschaltung") + 1  # Start just after the heading
            
            # Loop through paragraphs until the next major section or the end
            while i < len(para):
                # Check if a subheading like "Verschaltung" is present
                if para[i].startswith("Verschaltung"):
                    add_h3(doc, para[i])  # Add subheading

                    # Add the associated table
                    copy_table(doc, raw.tables[tableIndex])
                    print(f"Verschaltung {i} table copied - {tableIndex}")
                    tableIndex += 1
                    
                    # Check for another table under the same subheading
                    while i + 1 < len(para) and not para[i+1].startswith("Verschaltung") and not para[i+1] in h1:
                        i += 1
                        if para[i].startswith("Wechselrichter"):
                            # Process additional content if needed
                            # Assume the next table belongs to the same subheading
                            copy_table(doc, raw.tables[tableIndex])
                            print(f"Wechselrichter {i} table copied - {tableIndex}")
                            tableIndex += 1

                # Stop when reaching the next major section or the end of the document
                if i + 1 < len(para) and para[i+1] in h1:
                    break
                
                i += 1  # Move to the next paragraph

        if "AC-Netz" in h2 :
            add_h2(doc, "AC-Netz")
            i = para.index("AC-Netz")
            add_h3(doc, para[i+1])
            copy_table(doc, raw.tables[tableIndex])
            print(f"AC-Netz table copied - {tableIndex}")
            # # print(tableIndex)
            tableIndex+=1


        if "Batteriesysteme" in h2 :
            add_h2(doc, "Batteriesysteme")
            i = para.index("Batteriesysteme")
            for k in range(i+1, len(para), 1):
                if 'Batteriesystem' in para[k]:
                    add_h3(doc, para[k])
                    copy_table(doc, raw.tables[tableIndex])
                    print(f"Batteriesysteme table copied - {tableIndex}")
                    tableIndex+=1
                else:
                    break
          

        # ----------------------------------------
        doc.add_page_break()
        # ----------------------------------------
        try:
            add_h1(doc, f"3. {h1[h1_index]}")
            h1_index+=1
        except:
            pass
        if "Ergebnisse Gesamtanlage" in h2 :
            add_h2(doc, "Ergebnisse Gesamtanlage")
            i = para.index("Ergebnisse Gesamtanlage")
            i+=1
            if "PV-Anlage" in para:
                add_h3(doc, "PV-Anlage")
                # here index represent table no of the raw document
                copy_table(doc, raw.tables[tableIndex])
                print(f"PV-Anlage table copied - {tableIndex}")
                # print(tableIndex)
                # index represent table no of the output document
                try:
                    format_table_with_picture(doc, tableIndex-1, f"{folder_path}/images/{pic_index}.png")
                except:
                    format_table_with_picture(doc, tableIndex-1, f"{folder_path}/images/{pic_index}.jpg")
                pic_index+=1
           
                tableIndex+=1
                i+=1
            
            if "Verbraucher" in para:
                add_h3(doc, "Verbraucher")
                # here index represent table no of the raw document
                copy_table(doc, raw.tables[tableIndex])
                print(f"Verbraucher table copied - {tableIndex}")
                # print(tableIndex)
                if flag == 0:
                    pass
                else:
                    # index represent table no of the output document
                    try:
                        format_table_with_picture(doc, tableIndex-1, f"{folder_path}/images/{pic_index}.png")
                        
                    except: 
                        format_table_with_picture(doc, tableIndex-1, f"{folder_path}/images/{pic_index}.jpg")
                    pic_index+=1
                tableIndex+=1
                i+=1
            
            if "Batteriesystem" in para:
                add_h3(doc, "Batteriesystem")
                copy_table(doc, raw.tables[tableIndex])
                print(f"Batteriesystem table copied - {tableIndex}")
                # print(tableIndex)
                tableIndex+=1
                i+=1

            if "Autarkiegrad" in para:
                add_h3(doc, "Autarkiegrad")
                copy_table(doc, raw.tables[tableIndex])
                print(f"Autarkiegrad table copied - {tableIndex}")
                # print(tableIndex)
                tableIndex+=1
                i+=1

            doc.add_paragraph("")

            while True:
                if 'Abbildung' in para[i]: 
                    add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(6), height=Inches(4))
                    pic_index+=1
                    i+=1
                else: 
                    break

        if "Ergebnisse pro Modulfläche" in h2 :
            add_h2(doc, "Ergebnisse pro Modulfläche")
            i = para.index("Ergebnisse pro Modulfläche")
            while (para[i+1] != h1[h1_index]):
                try:
                    add_h3(doc, para[i+1])
                    copy_table(doc, raw.tables[tableIndex])
                    print(f"Ergebnisse pro Modulfläche {i} table copied - {tableIndex}")
                    # print(tableIndex)
                    tableIndex+=1
                    i+=1
                except:
                    break

        try:
            add_h1(doc, f"4. {h1[h1_index]}")
            h1_index+=1
        except:
            pass
        
        if "Energiebilanz Sankey-Diagramm" in h1 :
            add_h2(doc, "Energiebilanz Sankey-Diagramm")
            try:
                add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                pic_index+=1
                i+=1
            except:
                pass

        if "Datenblatt PV-Modul" in h2 :
            add_h2(doc, "Datenblatt PV-Modul")
            i = para.index("Datenblatt PV-Modul")
            for k in range(i+1, len(para), 1):
                if 'PV-Modul' in para[k]:
                    add_h3(doc, para[k])
                    copy_table(doc, raw.tables[tableIndex])
                    print(f"Datenblatt PV-Modul pro Modulfläche {k} table copied - {tableIndex}")
                    tableIndex+=1
                else:
                    break

        if "Datenblatt Wechselrichter" in h2 :
            add_h2(doc, "Datenblatt Wechselrichter")
            i = para.index("Datenblatt Wechselrichter")
            while (not para[i+1].startswith("Datenblatt Batteriesystem") and not para[i+1].startswith("Datenblatt Batterie") and not para[i+1].startswith("Schaltplan") and not para[i+1].startswith("Übersichtsplan") and not para[i+1].startswith("Bemaßungsplan") and not para[i+1].startswith("Strangplan") and not para[i+1].startswith("Stückliste")):
                try:
                    add_h3(doc, para[i+1])
                    copy_table(doc, raw.tables[tableIndex])
                    print(f"Datenblatt Wechselrichter table copied - {tableIndex}")
                    tableIndex+=1
                    i+=1
                except:
                    break

        
        if "Datenblatt Batteriesystem" in h2 :
            add_h2(doc, "Datenblatt Batteriesystem")
            i = para.index("Datenblatt Batteriesystem")
            while (para[i+1].startswith("Batteriesystem")):
                add_h3(doc, para[i+1])
                copy_table(doc, raw.tables[tableIndex])
                # print(tableIndex)
                tableIndex+=1
                i+=1


        if "Datenblatt Batterie" in h2 :
            add_h2(doc, "Datenblatt Batterie")
            i = para.index("Datenblatt Batterie")
            try:
                while (para[i+1] != h1[h1_index]):
                    add_h3(doc, para[i+1])
                    copy_table(doc, raw.tables[tableIndex])
                    # print(tableIndex)
                    tableIndex+=1
                    i+=1
            except:
                print(f'Total number of tables: {len(raw.tables)})')
                print(f'table index: {tableIndex}')

        # ----------------------------------------
        doc.add_page_break()
        # ----------------------------------------
        try:
            add_h1(doc, f"5. {h1[h1_index]}")
            h1_index+=1
        except:
            pass
        if "Schaltplan" in h2 :
            add_h2(doc, "Schaltplan")
            i = para.index("Schaltplan")
            while True:
                if i+1 < len(para) and 'Abbildung' in para[i+1]: 
                    add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                    pic_index+=1
                    i+=1
                else: 
                    break
        if "Übersichtsplan" in h2 :
            add_h2(doc, "Übersichtsplan")
            i = para.index("Übersichtsplan")
            while True:
                if i+1 < len(para) and 'Abbildung' in para[i+1]: 
                    add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                    pic_index+=1
                    i+=1
                else: 
                    break

        if "Bemaßungsplan" in h2 :
            add_h2(doc, "Bemaßungsplan")
            i = para.index("Bemaßungsplan")
            while True:
                if i+1 < len(para) and 'Abbildung' in para[i+1]: 
                    add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                    pic_index+=1
                    i+=1
                else: 
                    break
        doc.add_page_break()

        if "Strangplan" in h2 :
            add_h2(doc, "Strangplan")
            i = para.index("Strangplan")
            while True:
                if i+1 <= len(para) : 
                    add_picture_inline(doc, f"{folder_path}/images/{pic_index}", width=Inches(5.5), height=Inches(6.5))
                    pic_index+=1
                    i+=1
                else: 
                    break

        if "Stückliste" in h2 :
            add_h2(doc, "Stückliste")
            i = para.index("Stückliste")
            add_h3(doc, para[i+1])
            # print(tableIndex)
        if "Umgebung" in h2:
            add_h2(doc, "Umgebung")
            i = para.index("Umgebung")
            while True:
                if i+1 <= len(para) : 
                    add_picture_inline(doc, f"{folder_path}/images/{2}", width=Inches(5.5), height=Inches(3.5))
                    i+=1
                    break
                else: 
                    break
            
        # ----------------------------------------
        doc.add_page_break()
        # ----------------------------------------   
        
        add_h1(doc, "6. Warum Solardach24 GmbH?")
        doc.add_picture("assets/template_images/image2.png", width=Inches(6.5), height=Inches(7))

        #add_h1(doc, "7. Vier Köpfe. Für Ihre PV-Anlage.")
        #doc.add_picture("assets/template_images/image3.png", width=Inches(6.5), height=Inches(7))

        #add_h1(doc, "8. Unser Partner. Für Ihre Sicherheit.")
        #doc.add_picture("assets/template_images/image4.png", width=Inches(6.5), height=Inches(7))

        #add_h1(doc, "9. Unser Haustechnik-Partner. Für Ihre persönliche Energiewende.")
        add_h1(doc, "7. Wer wir sind.")
        doc.add_picture("assets/template_images/image3.png", width=Inches(6.5), height=Inches(7))

        add_h1(doc, "8. Unser Haustechnik-Partner. Für Ihre persönliche Energiewende.")
        doc.add_picture("assets/template_images/image4.png", width=Inches(5.8), height=Inches(5.5))

        add_h1(doc, "9. Unsere Elektropartner. Für Ihre Sicherheit.")
        doc.add_picture("assets/template_images/image5.png", width=Inches(6), height=Inches(6))

        add_h1(doc, "10. Unser Versicherungspartner. Exklusiv bei der Solardach24.")
        doc.add_picture("assets/template_images/image6.png", width=Inches(5), height=Inches(6))

        add_h1(doc, "11. Unsere Lieferanten. Für die besten Komponenten.")
        doc.add_picture("assets/template_images/image7.png", width=Inches(6), height=Inches(6))

        add_h1(doc, "12. Gesellschaftliches Engagement und Mitgliedschaften")
        doc.add_picture("assets/template_images/image8.png", width=Inches(6), height=Inches(6))

        prepare_header(doc, raw)
        prepare_footer(doc)
        replace_variables(doc, raw)
        set_font_to_barlow(doc)

        format_table(doc)
        darken_first_row_bottom_border(doc)
        update_module_name(doc)
        add_page_numbers(doc)  # Call the function here to add page numbers
        # Remove empty paragraphs and sections
        #remove_empty_paragraphs(doc)
        #remove_empty_sections(doc)
        doc.save(f'{output_folder}/{fileName}-output.docx')
        print(f'{fileName}-output.docx created')

        clear_folder_contents(fileName,folder_path)
        print("")
        print("")
        print("")
    except Exception as e:
        count += 1
        logging.error(f"\n\n{count}\nAn error occurred", exc_info=True)

        folder_path = os.path.dirname(filepath)
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path) and not filename.endswith('.docx'):
                os.remove(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
            elif filename.startswith('~$'):
                os.remove(file_path)
            elif filename.endswith('.docx') and filename == fileName:
                src_file = os.path.join(folder_path, filename)
                dst_file = os.path.join(INVALID_FOLDER, filename)
                shutil.move(src_file, dst_file)
                

if __name__=="__main__":
    
    #Deployment -- comment this section when testing
    #'''
    WATCH_FOLDER = 'D:/OneDrive/Office/OneDrive - Solardach24 GmbH/Intranet-Dokumente/10-Verkauf/00-Administration/DocGenerator/input/'
    TEMPLATE_PATH = 'assets/template.docx'
    OUTPUT_FOLDER = 'D:/OneDrive/Office/OneDrive - Solardach24 GmbH/Intranet-Dokumente/10-Verkauf/00-Administration/DocGenerator/output/'
    INVALID_FOLDER = 'D:/OneDrive/Office/OneDrive - Solardach24 GmbH/Intranet-Dokumente/10-Verkauf/00-Administration/DocGenerator/invalid/'
    #'''
    '''
    #Testing -- comment this section when testing   
    WATCH_FOLDER = 'input/'
    TEMPLATE_PATH = 'assets/template.docx'
    OUTPUT_FOLDER = 'output/'
    INVALID_FOLDER = 'invalid/'
    '''
    
    folders = [WATCH_FOLDER, OUTPUT_FOLDER, INVALID_FOLDER]
    # Check and create folders if needed
    for folder in folders:
        if not os.path.exists(folder):
            os.makedirs(folder)
            print(f'Created folder: {folder}')
        else:
            pass
    set_(WATCH_FOLDER, TEMPLATE_PATH, OUTPUT_FOLDER)

    